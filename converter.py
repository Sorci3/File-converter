
import sys
import os
import subprocess
from pathlib import Path


def out_path(src: Path, ext: str) -> Path:
    return src.with_suffix("." + ext)

def _download_ffmpeg(dest: Path):
    import urllib.request
    import zipfile
    import tempfile

    URL = (
        "https://github.com/BtbN/ffmpeg-builds/releases/download/latest/"
        "ffmpeg-master-latest-win64-gpl.zip"
    )
    print("ffmpeg not found — downloading automatically (~80 MB)...")
    dest.parent.mkdir(exist_ok=True)

    tmp = Path(tempfile.mktemp(suffix=".zip"))
    try:
        def _progress(block, block_size, total):
            if total > 0:
                print(f"\r  {min(block * block_size * 100 // total, 100)}%", end="", flush=True)

        urllib.request.urlretrieve(URL, tmp, _progress)
        print("\nExtracting...")

        with zipfile.ZipFile(tmp) as zf:
            entry = next((n for n in zf.namelist() if n.endswith("/bin/ffmpeg.exe")), None)
            if not entry:
                raise FileNotFoundError("ffmpeg.exe not found in archive.")
            dest.write_bytes(zf.read(entry))

        print(f"ffmpeg installed: {dest}")
    finally:
        tmp.unlink(missing_ok=True)

def find_ffmpeg() -> str:
    local = Path(__file__).parent / "ffmpeg" / "ffmpeg.exe"
    if local.exists():
        return str(local)
    from shutil import which
    path = which("ffmpeg")
    if path:
        return path
    _download_ffmpeg(local)
    return str(local)

# --- Images ---

IMAGE_EXTS = {"jpg", "jpeg", "png", "gif", "bmp", "tiff", "tif", "webp"}
IMAGE_FORMATS = {"jpeg": "JPEG", "jpg": "JPEG", "png": "PNG",
                 "gif": "GIF", "bmp": "BMP", "tiff": "TIFF", "tif": "TIFF", "webp": "WEBP"}

def convert_image(src: Path, fmt: str):
    from PIL import Image
    pil_fmt = IMAGE_FORMATS.get(fmt, fmt.upper())
    dst = out_path(src, "jpg" if fmt == "jpeg" else fmt)
    img = Image.open(src)
    # RGBA → RGB required for JPEG (no transparency support)
    if pil_fmt == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(dst, pil_fmt)
    print(f"Done: {dst}")

# --- Audio / Video ---

MEDIA_EXTS = {"mp4", "avi", "mov", "wmv", "flv", "mp3", "wav", "aac", "flac", "ogg"}

def convert_media(src: Path, fmt: str):
    ffmpeg = find_ffmpeg()
    dst = out_path(src, fmt)
    result = subprocess.run([ffmpeg, "-i", str(src), "-y", str(dst)], capture_output=True, text=True)
    if result.returncode != 0:
        print(f"ffmpeg error:\n{result.stderr[-500:]}", file=sys.stderr)
        sys.exit(1)
    print(f"Done: {dst}")

# --- Documents ---

DOC_EXTS = {"pdf", "docx", "txt"}

def txt_to_docx(src: Path):
    from docx import Document
    doc = Document()
    doc.add_paragraph(src.read_text(encoding="utf-8", errors="replace"))
    dst = out_path(src, "docx")
    doc.save(dst)
    print(f"Done: {dst}")

def txt_to_pdf(src: Path):
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for line in src.read_text(encoding="utf-8", errors="replace").splitlines():
        pdf.cell(0, 6, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    dst = out_path(src, "pdf")
    pdf.output(str(dst))
    print(f"Done: {dst}")

def docx_to_txt(src: Path):
    from docx import Document
    doc = Document(src)
    dst = out_path(src, "txt")
    dst.write_text("\n".join(p.text for p in doc.paragraphs), encoding="utf-8")
    print(f"Done: {dst}")

def docx_to_pdf(src: Path):
    try:
        import win32com.client
    except ImportError:
        print("Error: pip install pywin32 required for DOCX→PDF.", file=sys.stderr)
        sys.exit(1)
    dst = out_path(src, "pdf")
    word = win32com.client.Dispatch("Word.Application")
    try:
        doc = word.Documents.Open(str(src.resolve()))
        doc.SaveAs(str(dst.resolve()), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
    finally:
        word.Quit()
    print(f"Done: {dst}")

def pdf_to_txt(src: Path):
    from pypdf import PdfReader
    text = "\n".join(page.extract_text() or "" for page in PdfReader(src).pages)
    dst = out_path(src, "txt")
    dst.write_text(text, encoding="utf-8")
    print(f"Done: {dst}")

def convert_doc(src: Path, fmt: str):
    ext = src.suffix.lstrip(".").lower()
    routes = {
        ("txt",  "docx"): lambda: txt_to_docx(src),
        ("txt",  "pdf"):  lambda: txt_to_pdf(src),
        ("docx", "txt"):  lambda: docx_to_txt(src),
        ("docx", "pdf"):  lambda: docx_to_pdf(src),
        ("pdf",  "txt"):  lambda: pdf_to_txt(src),
    }
    fn = routes.get((ext, fmt))
    if fn is None:
        print(f"Error: {ext}→{fmt} conversion not supported.", file=sys.stderr)
        sys.exit(1)
    fn()

# --- Main ---

def main():
    if len(sys.argv) != 3:
        print("Usage: converter.py <file> <format>")
        sys.exit(1)

    src = Path(sys.argv[1])
    fmt = sys.argv[2].lower().lstrip(".")

    if not src.exists():
        print(f"File not found: {src}", file=sys.stderr)
        sys.exit(1)

    ext = src.suffix.lstrip(".").lower()

    if ext in IMAGE_EXTS:
        convert_image(src, fmt)
    elif ext in MEDIA_EXTS:
        convert_media(src, fmt)
    elif ext in DOC_EXTS:
        convert_doc(src, fmt)
    else:
        print(f"Error: .{ext} format not supported.", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
